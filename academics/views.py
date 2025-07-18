# D:\classroom\academics\views.py

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
from django.db import transaction, IntegrityError
from decimal import Decimal
import openpyxl
import docx

from students.models import Student
from .models import Classroom, Enrollment, Assignment, Submission
from .forms import AssignmentForm
from teachers.forms import StudentForm, UploadFileForm

# ===================================================================
# Views จัดการ Assignment
# ===================================================================
@login_required
def assignment_detail_view(request, pk):
    # ... โค้ด ...
    assignment = get_object_or_404(Assignment, pk=pk, classroom__subject__teacher=request.user)
    # ในฟังก์ชัน assignment_detail_view
    students = Student.objects.filter(enrollments__classroom=assignment.classroom).order_by('student_id')
    submissions = Submission.objects.filter(assignment=assignment).select_related('enrollment__student')
    submission_dict = {sub.enrollment.student.pk: sub for sub in submissions}
    context = {
        'assignment': assignment,
        'all_students': students,
        'submission_dict': submission_dict,
    }
    return render(request, 'academics/assignment_detail.html', context)

@login_required
def assignment_update_view(request, pk):
    # ... โค้ด ...
    assignment = get_object_or_404(Assignment, pk=pk, classroom__subject__teacher=request.user)
    if request.method == 'POST':
        form = AssignmentForm(request.POST, request.FILES, instance=assignment)
        if form.is_valid():
            form.save()
            messages.success(request, f"แก้ไขชิ้นงาน '{assignment.title}' เรียบร้อยแล้ว")
            return redirect('academics:assignment_detail', pk=assignment.pk)
    else:
        form = AssignmentForm(instance=assignment)
    context = {'form': form, 'assignment': assignment}
    return render(request, 'academics/assignment_form.html', context)


@login_required
def assignment_delete_view(request, pk):
    # ... โค้ด ...
    assignment = get_object_or_404(Assignment, pk=pk, classroom__subject__teacher=request.user)
    if request.method == 'POST':
        classroom_pk = assignment.classroom.pk
        assignment_title = assignment.title
        assignment.delete()
        messages.success(request, f"ลบชิ้นงาน '{assignment_title}' เรียบร้อยแล้ว")
        return redirect('teachers:classroom_detail', pk=classroom_pk)
    context = {'assignment': assignment}
    return render(request, 'academics/assignment_confirm_delete.html', context)

# ===================================================================
# Views จัดการ Submission และคะแนน
# ===================================================================
@login_required
def grade_submission_view(request, pk):
    # ... โค้ด ...
    submission = get_object_or_404(Submission, pk=pk, assignment__classroom__subject__teacher=request.user)
    if request.method == 'POST':
        grade_str = request.POST.get('grade', '').strip()
        feedback = request.POST.get('feedback', '').strip()

        if grade_str:
            try:
                submission.grade = Decimal(grade_str)
            except (ValueError, TypeError):
                messages.error(request, "กรุณากรอกคะแนนเป็นตัวเลขที่ถูกต้อง")
                return redirect('academics:assignment_detail', pk=submission.assignment.pk)
        else:
            submission.grade = None

        submission.feedback = feedback
        submission.save()
        messages.success(request, f"บันทึกคะแนนและคอมเมนต์ของ {submission.student.full_name} เรียบร้อยแล้ว")
    return redirect('academics:assignment_detail', pk=submission.assignment.pk)

@login_required
def gradebook_view(request, pk):
    # ... โค้ด ...
    classroom = get_object_or_404(Classroom, pk=pk, subject__teacher=request.user)
    enrollments = Enrollment.objects.filter(classroom=classroom).select_related('student').order_by('student__student_id')
    assignments = classroom.assignments.all().order_by('created_at')
    submissions = Submission.objects.filter(assignment__in=assignments).select_related('enrollment')
    submission_map = {(sub.enrollment_id, sub.assignment_id): sub.grade for sub in submissions}
    gradebook_data = []
    for enrollment in enrollments:
        scores = {
            assignment.id: submission_map.get((enrollment.id, assignment.id))
            for assignment in assignments
        }
        gradebook_data.append({
            'student': enrollment.student,
            'enrollment': enrollment,
            'scores': scores
        })
    context = {
        'classroom': classroom,
        'assignments': assignments,
        'gradebook_data': gradebook_data,
    }
    return render(request, 'academics/gradebook.html', context)

# ===================================================================
# Views จัดการ Student (ที่ถูกเรียกใช้โดย Teacher)
# ===================================================================
@login_required
def student_update_view(request, pk):
    # ... โค้ด ...
    student = get_object_or_404(Student, pk=pk)
    if not Enrollment.objects.filter(student=student, classroom__subject__teacher=request.user).exists():
        messages.error(request, "คุณไม่มีสิทธิ์แก้ไขข้อมูลนักเรียนคนนี้")
        return redirect('teachers:dashboard')

    if request.method == 'POST':
        form = StudentForm(request.POST, instance=student)
        if form.is_valid():
            form.save()
            messages.success(request, f"อัปเดตข้อมูลนักเรียน '{student.full_name}' สำเร็จ")
            return redirect(request.GET.get('next', 'teachers:dashboard'))
    else:
        form = StudentForm(instance=student)
    context = {'form': form, 'student': student}
    return render(request, 'academics/student_update_form.html', context)

@login_required
def student_remove_from_classroom_view(request, classroom_pk, student_pk):
    # ... โค้ด ...
    enrollment = get_object_or_404(Enrollment, classroom_id=classroom_pk, student_id=student_pk, classroom__subject__teacher=request.user)
    student_name = enrollment.student.full_name
    if request.method == 'POST':
        enrollment.delete()
        messages.success(request, f"นำนักเรียน '{student_name}' ออกจากห้องเรียนเรียบร้อยแล้ว")
        return redirect('teachers:classroom_detail', pk=classroom_pk)
    context = {'enrollment': enrollment}
    return render(request, 'academics/student_confirm_remove.html', context)

# ===================================================================
# Views สำหรับ Import/Export และเครื่องมืออื่นๆ
# ===================================================================
@login_required
def upload_students_view(request):
    # ... โค้ด ...
    if request.method != 'POST':
        return redirect('teachers:dashboard')
    form = UploadFileForm(data=request.POST, files=request.FILES, user=request.user)
    if not form.is_valid():
        for field, errors in form.errors.items():
            for error in errors:
                messages.error(request, f"ฟอร์มผิดพลาด ({field}): {error}")
        return redirect(request.META.get('HTTP_REFERER', 'teachers:dashboard'))
    classroom = form.cleaned_data['classroom']
    uploaded_file = request.FILES['file']
    file_extension = uploaded_file.name.split('.')[-1].lower()
    try:
        if file_extension == 'xlsx':
            added, updated, errors = process_uploaded_file(uploaded_file, classroom, is_word=False)
        elif file_extension in ['doc', 'docx']:
            added, updated, errors = process_uploaded_file(uploaded_file, classroom, is_word=True)
        else:
            messages.error(request, "ไม่รองรับไฟล์ประเภทนี้ กรุณาใช้ไฟล์ .xlsx หรือ .docx")
            return redirect('teachers:classroom_detail', pk=classroom.pk)
        if added > 0 or updated > 0:
            success_message = f"นำเข้าสำเร็จ! เพิ่มนักเรียนใหม่ {added} คน, พบนักเรียนที่มีอยู่แล้ว {updated} คน"
            messages.success(request, success_message)
        if errors > 0:
            messages.warning(request, f"มี {errors} แถวในไฟล์ที่ไม่สามารถประมวลผลได้ (ข้อมูลอาจไม่ครบ)")
        if added == 0 and updated == 0 and errors > 0:
             messages.error(request, "ไม่สามารถนำเข้าข้อมูลจากไฟล์ได้ กรุณาตรวจสอบรูปแบบไฟล์")
    except Exception as e:
        messages.error(request, f"เกิดข้อผิดพลาดร้ายแรงในการประมวลผลไฟล์: {e}")
    return redirect('teachers:classroom_detail', pk=classroom.pk)

def process_uploaded_file(file, classroom, is_word=False):
    # ... โค้ด ...
    added_count, updated_count, error_count = 0, 0, 0
    all_rows = []
    try:
        if is_word:
            document = docx.Document(file)
            if not document.tables: return 0, 0, 1
            table = document.tables[0]
            all_rows = [tuple(cell.text.strip() for cell in row.cells) for row in table.rows]
        else:
            workbook = openpyxl.load_workbook(file, data_only=True)
            sheet = workbook.active
            all_rows = list(sheet.iter_rows(values_only=True))
    except Exception as e:
        print(f"Error reading file: {e}")
        return 0, 0, 1
    if not all_rows: return 0, 0, 0
    start_row_index = 0
    first_row = all_rows[0]
    try:
        if first_row and first_row[0]: int(str(first_row[0]).strip())
    except (ValueError, TypeError, IndexError):
        start_row_index = 1
    for row_data in all_rows[start_row_index:]:
        if len(row_data) < 2 or not row_data[0] or not str(row_data[0]).strip() or not row_data[1] or not str(row_data[1]).strip():
            error_count += 1
            continue
        student_id = str(row_data[0]).strip()
        full_name = str(row_data[1]).strip()
        try:
            with transaction.atomic():
                student, student_created = Student.objects.get_or_create(student_id=student_id, defaults={'full_name': full_name})
                if student_created or not student.user:
                    user, user_created = User.objects.get_or_create(username=student_id)
                    if user_created:
                        name_parts = full_name.split(' ', 1)
                        user.first_name = name_parts[0]
                        user.last_name = name_parts[1] if len(name_parts) > 1 else ''
                        user.set_password('qw123456')
                        user.save()
                    student.user = user
                    student.save()
                _enrollment, enrolled_now = Enrollment.objects.get_or_create(student=student, classroom=classroom)
                if enrolled_now:
                    added_count += 1
                else:
                    updated_count += 1
        except Exception as e:
            print(f"Error processing row (ID: {student_id}): {e}")
            error_count += 1
    return added_count, updated_count, error_count

@login_required
def copy_students_view(request, to_classroom_pk):
    # ... โค้ด ...
    if request.method == 'POST':
        from_classroom_pk = request.POST.get('from_classroom')
        try:
            to_classroom = get_object_or_404(Classroom, pk=to_classroom_pk, subject__teacher=request.user)
            from_classroom = get_object_or_404(Classroom, pk=from_classroom_pk, subject__teacher=request.user)
            students_to_copy = from_classroom.students.all()
            if not students_to_copy.exists():
                messages.warning(request, f"ห้องเรียนต้นทาง '{from_classroom}' ไม่มีนักเรียนให้คัดลอก")
                return redirect('teachers:classroom_detail', pk=to_classroom_pk)
            copied_count = 0
            for student in students_to_copy:
                _enrollment, created = Enrollment.objects.get_or_create(student=student, classroom=to_classroom)
                if created:
                    copied_count += 1
            messages.success(request, f"คัดลอกนักเรียน {copied_count} คนจากห้อง '{from_classroom}' มายัง '{to_classroom}' เรียบร้อยแล้ว")
        except Classroom.DoesNotExist:
            messages.error(request, "ไม่พบห้องเรียนที่ระบุ หรือคุณไม่มีสิทธิ์เข้าถึง")
        except Exception as e:
            messages.error(request, f"เกิดข้อผิดพลาด: {e}")
    return redirect('teachers:classroom_detail', pk=to_classroom_pk)